import io
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document

from src.exceptions import ValidationError
from src.interfaces.content_extractor import IContentExtractor

DOCX_MAGIC = b"PK\x03\x04"
DOC_MAGIC = b"\xd0\xcf\x11\xe0"

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_TAG_P = f"{{{_W_NS}}}p"
_TAG_TBL = f"{{{_W_NS}}}tbl"
_TAG_TR = f"{{{_W_NS}}}tr"
_TAG_TC = f"{{{_W_NS}}}tc"
_TAG_T = f"{{{_W_NS}}}t"
_TAG_TAB = f"{{{_W_NS}}}tab"
_TAG_BR = f"{{{_W_NS}}}br"
_TAG_LRPB = f"{{{_W_NS}}}lastRenderedPageBreak"
_ATTR_TYPE = f"{{{_W_NS}}}type"


class WordExtractor(IContentExtractor):
    def from_bytes(self, data: bytes) -> str:
        if data.startswith(DOCX_MAGIC):
            return self._extract_docx(data)
        if data.startswith(DOC_MAGIC):
            return self._extract_doc(data)
        raise ValidationError("Word format not recognized. Expected .doc or .docx.")

    def pages_from_bytes(self, data: bytes) -> list[str] | None:
        if not data.startswith(DOCX_MAGIC):
            return None
        pages = self._extract_docx_pages(data)
        if len(pages) <= 1:
            return None
        return pages

    @staticmethod
    def _extract_docx(data: bytes) -> str:
        document = Document(io.BytesIO(data))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()

    @staticmethod
    def _extract_docx_pages(data: bytes) -> list[str]:
        document = Document(io.BytesIO(data))
        pages: list[list[str]] = [[]]

        def emit(text: str) -> None:
            if text:
                pages[-1].append(text)

        def new_page() -> None:
            pages.append([])

        body = document.element.body
        for block in body.iterchildren():
            tag = block.tag
            if tag == _TAG_P:
                for elem in block.iter():
                    et = elem.tag
                    if et == _TAG_LRPB or et == _TAG_BR and elem.get(_ATTR_TYPE) == "page":
                        new_page()
                    elif et == _TAG_T:
                        emit(elem.text or "")
                    elif et == _TAG_TAB:
                        emit("\t")
                emit("\n")
            elif tag == _TAG_TBL:
                for row in block.iter(_TAG_TR):
                    cells: list[str] = []
                    for cell in row.iter(_TAG_TC):
                        parts: list[str] = []
                        for t in cell.iter(_TAG_T):
                            if t.text:
                                parts.append(t.text)
                        cells.append("".join(parts).strip())
                    if any(cells):
                        emit(" | ".join(cells) + "\n")

        result = ["".join(parts).strip() for parts in pages]
        while result and not result[-1]:
            result.pop()
        return result or [""]

    @staticmethod
    def _extract_doc(data: bytes) -> str:
        if shutil.which("antiword") is None:
            raise ValidationError(
                "Support for .doc requires the 'antiword' utility installed on the system."
            )
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        try:
            result = subprocess.run(
                ["antiword", tmp_path],
                capture_output=True,
                text=True,
                check=True,
            )
            return result.stdout.strip()
        except subprocess.CalledProcessError as exc:
            raise ValidationError(
                f"Failed to extract text from .doc: {exc.stderr.strip() or exc}"
            ) from exc
        finally:
            Path(tmp_path).unlink(missing_ok=True)
