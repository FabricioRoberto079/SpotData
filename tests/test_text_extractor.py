import pytest

from src.enums.content_type import ContentType
from src.exceptions import ValidationError
from src.interfaces.content_extractor import IContentExtractor
from src.services.text_extractor import TextExtractor


class _Stub(IContentExtractor):
    def __init__(self, value: str):
        self.value = value
        self.from_bytes_calls: list[bytes] = []

    def from_bytes(self, data):
        self.from_bytes_calls.append(data)
        return self.value


def test_dispatches_per_content_type():
    pdf = _Stub("PDF")
    txt = _Stub("TXT")
    extractor = TextExtractor({ContentType.PDF: pdf, ContentType.TEXTO: txt})
    assert extractor.extract_from_bytes(b"x", ContentType.PDF) == "PDF"
    assert extractor.extract_from_bytes(b"x", ContentType.TEXTO) == "TXT"
    assert pdf.from_bytes_calls == [b"x"]


def test_unsupported_type_raises_validation_error():
    extractor = TextExtractor({ContentType.PDF: _Stub("x")})
    with pytest.raises(ValidationError):
        extractor.extract_from_bytes(b"x", ContentType.FOTO)


def test_plain_text_extractor_decodes_utf8():
    from src.services.extractors.plain_text import PlainTextExtractor

    assert PlainTextExtractor().from_bytes("olá".encode()) == "olá"


def test_word_extractor_splits_docx_on_manual_page_break():
    import io

    from docx import Document
    from docx.enum.text import WD_BREAK

    from src.services.extractors.word import WordExtractor

    doc = Document()
    doc.add_paragraph("primeira página")
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("segunda página")
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("terceira página")

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    pages = WordExtractor().pages_from_bytes(data)
    assert pages is not None
    assert len(pages) == 3
    assert "primeira página" in pages[0]
    assert "segunda página" in pages[1]
    assert "terceira página" in pages[2]


def test_word_extractor_returns_none_for_single_page_docx():
    import io

    from docx import Document

    from src.services.extractors.word import WordExtractor

    doc = Document()
    doc.add_paragraph("apenas uma página, sem quebras")
    buf = io.BytesIO()
    doc.save(buf)

    assert WordExtractor().pages_from_bytes(buf.getvalue()) is None
